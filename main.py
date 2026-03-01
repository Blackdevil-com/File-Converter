import signal
import os
import uuid
import shutil
import threading
import time
import zipfile
import subprocess

from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import FileResponse

import fitz
from PyPDF2 import PdfMerger, PdfReader, PdfWriter
from docx import Document
import pytesseract
from pdf2image import convert_from_path

import uvicorn
from converter import (
    convert_multiple_images_to_pdf,
    convert_pdf_to_word,
    convert_to_pdf,
    create_text_watermark,
    excel_to_pdf,
    pdf_extract_images,
    pdf_pages_to_images
)

app = FastAPI()

UPLOAD_DIR = "/home/ubuntu/uploads"
OUTPUT_DIR = "/home/ubuntu/outputs"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ---------------- Helper Functions ----------------
def delete_later(path, delay=3600):
    """Delete a file after `delay` seconds."""
    time.sleep(delay)
    if os.path.exists(path):
        os.remove(path)

def get_unique_filename(folder, filename):
    """Return a unique filename if the file exists in folder."""
    base, ext = os.path.splitext(filename)
    counter = 1
    new_filename = filename
    while os.path.exists(os.path.join(folder, new_filename)):
        new_filename = f"{base} ({counter}){ext}"
        counter += 1
    return os.path.join(folder, new_filename)


# ---------------- Shutdown Endpoint ----------------
@app.get("/shutdown")
def shutdown_server():
    os.kill(os.getpid(), signal.SIGTERM)
    return {"message": "Server is shutting down"}


# ---------------- Word to PDF ----------------
@app.post("/convert/word-to-pdf")
async def word_to_pdf(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".docx"):
        return {"error": "Only .docx files allowed"}

    unique_id = str(uuid.uuid4())
    input_path = os.path.join(UPLOAD_DIR, unique_id + ".docx")
    output_name = os.path.splitext(file.filename)[0] + ".pdf"
    output_path = get_unique_filename(OUTPUT_DIR, output_name)

    with open(input_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    try:
        convert_to_pdf(input_path, OUTPUT_DIR)
        generated_pdf = os.path.join(OUTPUT_DIR, unique_id + ".pdf")
        os.rename(generated_pdf, output_path)
    except Exception as e:
        return {"error": str(e)}

    threading.Thread(target=delete_later, args=(input_path,)).start()
    threading.Thread(target=delete_later, args=(output_path,)).start()

    return FileResponse(output_path, filename=output_name)


# ---------------- PDF to Word ----------------
@app.post("/convert/pdf-to-word")
async def pdf_to_word(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".pdf"):
        return {"error": "Only .pdf files allowed"}

    unique_id = str(uuid.uuid4())
    input_path = os.path.join(UPLOAD_DIR, unique_id + ".pdf")
    output_name = os.path.splitext(file.filename)[0] + ".docx"
    output_path = get_unique_filename(OUTPUT_DIR, output_name)

    with open(input_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    try:
        convert_pdf_to_word(input_path, output_path)
    except Exception as e:
        return {"error": str(e)}

    threading.Thread(target=delete_later, args=(input_path,)).start()
    threading.Thread(target=delete_later, args=(output_path,)).start()

    return FileResponse(
        output_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=output_name
    )

# ---------------- Multi-Image to PDF ----------------
@app.post("/convert/images-to-pdf")
async def multi_image_to_pdf(
    files: list[UploadFile] = File(...),
    orientation: str = Form("portrait"),
    scale: str = Form("full")
):
    if not files:
        return {"error": "No files uploaded"}

    allowed_ext = [".jpg", ".jpeg", ".png", ".webp"]
    input_paths = []
    unique_id = str(uuid.uuid4())

    # Save uploaded images internally
    for idx, file in enumerate(files):
        ext = os.path.splitext(file.filename)[1].lower()
        if ext not in allowed_ext:
            return {"error": f"Invalid file type: {file.filename}"}

        # Internal safe name (UUID + index)
        image_filename = f"{unique_id}_{idx}{ext}"
        image_path = os.path.join(UPLOAD_DIR, image_filename)
        contents = await file.read()
        with open(image_path, "wb") as f:
            f.write(contents)
        input_paths.append(image_path)

    # Internal output path (always unique with UUID)
    output_path = os.path.join(OUTPUT_DIR, f"{unique_id}.pdf")

    # Convert images to PDF
    try:
        convert_multiple_images_to_pdf(
            input_paths,
            output_path,
            orientation=orientation,
            scale=scale
        )
    except Exception as e:
        return {"error": f"Image to PDF conversion failed: {str(e)}"}

    # Cleanup input images in background
    for path in input_paths:
        threading.Thread(target=delete_later, args=(path,)).start()
    threading.Thread(target=delete_later, args=(output_path,)).start()

    # Unique name for **user download**
    user_filename = get_unique_filename(OUTPUT_DIR, "images.pdf")  # browser sees unique filename each time

    return FileResponse(
        output_path,
        media_type="application/pdf",
        filename=user_filename  # user will download as images_<uuid>.pdf
    )

# ---------------- PDF to Images ----------------
@app.post("/convert/pdf-image")
async def pdf_to_image_mode(file: UploadFile = File(...), mode: str = Form("page")):
    if not file.filename.lower().endswith(".pdf"):
        return {"error": "Only PDF allowed"}

    unique_id = str(uuid.uuid4())
    input_path = os.path.join(UPLOAD_DIR, unique_id + ".pdf")
    output_folder = os.path.join(OUTPUT_DIR, unique_id)
    os.makedirs(output_folder, exist_ok=True)

    contents = await file.read()
    with open(input_path, "wb") as f:
        f.write(contents)

    try:
        if mode == "page":
            image_paths = pdf_pages_to_images(input_path, output_folder)
        elif mode == "extract":
            image_paths = pdf_extract_images(input_path, output_folder)
            if not image_paths:
                image_paths = pdf_pages_to_images(input_path, output_folder)
        else:
            return {"error": "Invalid mode, choose 'page' or 'extract'"}
    except Exception as e:
        return {"error": str(e)}

    zip_path = get_unique_filename(OUTPUT_DIR, f"{os.path.splitext(file.filename)[0]}_images.zip")
    with zipfile.ZipFile(zip_path, "w") as zipf:
        for img in image_paths:
            zipf.write(img, os.path.basename(img))

    threading.Thread(target=delete_later, args=(input_path,)).start()
    threading.Thread(target=delete_later, args=(zip_path,)).start()

    return FileResponse(zip_path, media_type="application/zip", filename=os.path.basename(zip_path))


# ---------------- PDF Merge ----------------
@app.post("/convert/merge-pdfs")
async def merge_pdfs(files: list[UploadFile] = File(...)):
    if len(files) < 2:
        return {"error": "Upload at least 2 PDFs"}

    input_paths = []
    unique_id = str(uuid.uuid4())

    for file in files:
        if not file.filename.lower().endswith(".pdf"):
            return {"error": f"Invalid file type: {file.filename}"}
        path = os.path.join(UPLOAD_DIR, f"{uuid.uuid4()}.pdf")
        contents = await file.read()
        with open(path, "wb") as f:
            f.write(contents)
        input_paths.append(path)

    output_path = get_unique_filename(OUTPUT_DIR, f"merged_{unique_id}.pdf")
    try:
        merger = PdfMerger()
        for path in input_paths:
            merger.append(path)
        merger.write(output_path)
        merger.close()
    except Exception as e:
        return {"error": str(e)}

    for path in input_paths:
        threading.Thread(target=delete_later, args=(path,)).start()
    threading.Thread(target=delete_later, args=(output_path,)).start()

    return FileResponse(output_path, media_type="application/pdf", filename=os.path.basename(output_path))


# ---------------- PDF Split ----------------
@app.post("/convert/split-pdf")
async def split_pdf(file: UploadFile = File(...), start_page: int = Form(...), end_page: int = Form(...)):
    if not file.filename.lower().endswith(".pdf"):
        return {"error": "Only PDF files allowed"}

    unique_id = str(uuid.uuid4())
    input_path = os.path.join(UPLOAD_DIR, unique_id + ".pdf")
    output_path = get_unique_filename(OUTPUT_DIR, f"split_{unique_id}.pdf")

    contents = await file.read()
    with open(input_path, "wb") as f:
        f.write(contents)

    try:
        reader = PdfReader(input_path)
        writer = PdfWriter()
        total_pages = len(reader.pages)
        if start_page < 1 or end_page > total_pages or start_page > end_page:
            return {"error": f"Invalid page range. Total pages: {total_pages}"}

        for i in range(start_page - 1, end_page):
            writer.add_page(reader.pages[i])
        with open(output_path, "wb") as f:
            writer.write(f)
    except Exception as e:
        return {"error": str(e)}

    threading.Thread(target=delete_later, args=(input_path,)).start()
    threading.Thread(target=delete_later, args=(output_path,)).start()

    return FileResponse(output_path, media_type="application/pdf", filename=os.path.basename(output_path))


# ---------------- PDF Rotate ----------------
@app.post("/convert/rotate-pdf")
async def rotate_pdf(file: UploadFile = File(...), angle: int = Form(...), pages: str = Form("all")):
    if not file.filename.lower().endswith(".pdf"):
        return {"error": "Only PDF files allowed"}
    if angle not in [90, 180, 270]:
        return {"error": "Angle must be 90, 180, or 270"}

    unique_id = str(uuid.uuid4())
    input_path = os.path.join(UPLOAD_DIR, unique_id + ".pdf")
    output_path = get_unique_filename(OUTPUT_DIR, f"rotated_{unique_id}.pdf")

    contents = await file.read()
    with open(input_path, "wb") as f:
        f.write(contents)

    try:
        reader = PdfReader(input_path)
        writer = PdfWriter()
        total_pages = len(reader.pages)
        if pages.lower() == "all":
            page_indexes = range(total_pages)
        else:
            page_indexes = [int(p.strip()) - 1 for p in pages.split(",") if 0 <= int(p.strip()) - 1 < total_pages]

        for i in range(total_pages):
            page = reader.pages[i]
            if i in page_indexes:
                page.rotate(angle)
            writer.add_page(page)
        with open(output_path, "wb") as f:
            writer.write(f)
    except Exception as e:
        return {"error": str(e)}

    threading.Thread(target=delete_later, args=(input_path,)).start()
    threading.Thread(target=delete_later, args=(output_path,)).start()

    return FileResponse(output_path, media_type="application/pdf", filename=os.path.basename(output_path))


# ---------------- PDF Watermark ----------------
@app.post("/convert/watermark-pdf")
async def watermark_pdf(file: UploadFile = File(...), watermark_text: str = Form(...)):
    if not file.filename.lower().endswith(".pdf"):
        return {"error": "Only PDF files allowed"}

    unique_id = str(uuid.uuid4())
    input_path = os.path.join(UPLOAD_DIR, unique_id + ".pdf")
    watermark_path = os.path.join(OUTPUT_DIR, f"wm_{unique_id}.pdf")
    output_path = get_unique_filename(OUTPUT_DIR, f"watermarked_{unique_id}.pdf")

    contents = await file.read()
    with open(input_path, "wb") as f:
        f.write(contents)

    try:
        create_text_watermark(watermark_text, watermark_path)
        reader = PdfReader(input_path)
        watermark = PdfReader(watermark_path).pages[0]
        writer = PdfWriter()
        for page in reader.pages:
            page.merge_page(watermark)
            writer.add_page(page)
        with open(output_path, "wb") as f:
            writer.write(f)
    except Exception as e:
        return {"error": str(e)}

    threading.Thread(target=delete_later, args=(input_path,)).start()
    threading.Thread(target=delete_later, args=(watermark_path,)).start()
    threading.Thread(target=delete_later, args=(output_path,)).start()

    return FileResponse(output_path, media_type="application/pdf", filename=os.path.basename(output_path))


# ---------------- PDF Extract Text ----------------
@app.post("/convert/extract-text-pdf")
async def extract_pdf_text(file: UploadFile = File(...), output_format: str = Form("txt")):
    if not file.filename.lower().endswith(".pdf"):
        return {"error": "Only PDF files allowed"}

    unique_id = str(uuid.uuid4())
    input_path = os.path.join(UPLOAD_DIR, unique_id + ".pdf")
    output_path = get_unique_filename(OUTPUT_DIR, f"{unique_id}.{output_format}")

    contents = await file.read()
    with open(input_path, "wb") as f:
        f.write(contents)

    try:
        doc = fitz.open(input_path)
        all_text = "\n".join([page.get_text() for page in doc])
        if output_format == "txt":
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(all_text)
        elif output_format == "docx":
            docx_file = Document()
            for line in all_text.split("\n"):
                docx_file.add_paragraph(line)
            docx_file.save(output_path)
        else:
            return {"error": "Invalid output format"}
    except Exception as e:
        return {"error": str(e)}

    threading.Thread(target=delete_later, args=(input_path,)).start()
    threading.Thread(target=delete_later, args=(output_path,)).start()

    return FileResponse(output_path, media_type="application/octet-stream", filename=os.path.basename(output_path))


# ---------------- PDF OCR ----------------
@app.post("/convert/ocr-pdf")
async def ocr_pdf(file: UploadFile = File(...), output_format: str = Form("txt")):
    if not file.filename.lower().endswith(".pdf"):
        return {"error": "Only PDF files allowed"}

    unique_id = str(uuid.uuid4())
    input_path = os.path.join(UPLOAD_DIR, unique_id + ".pdf")
    output_path = get_unique_filename(OUTPUT_DIR, f"{unique_id}.{output_format}")

    contents = await file.read()
    with open(input_path, "wb") as f:
        f.write(contents)

    try:
        pages = convert_from_path(input_path, dpi=300)
        all_text = ""
        for img in pages:
            all_text += pytesseract.image_to_string(img) + "\n"
        if output_format == "txt":
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(all_text)
        elif output_format == "docx":
            docx_file = Document()
            for line in all_text.split("\n"):
                docx_file.add_paragraph(line)
            docx_file.save(output_path)
        else:
            return {"error": "Invalid output format"}
    except Exception as e:
        return {"error": str(e)}

    threading.Thread(target=delete_later, args=(input_path,)).start()
    threading.Thread(target=delete_later, args=(output_path,)).start()

    return FileResponse(output_path, media_type="application/octet-stream", filename=os.path.basename(output_path))


# ---------------- Excel to PDF ----------------
@app.post("/convert/excel-to-pdf")
async def convert_excel(file: UploadFile = File(...)):
    if not file.filename.lower().endswith((".xlsx", ".xls")):
        return {"error": "Only Excel files allowed"}

    unique_id = str(uuid.uuid4())
    ext = os.path.splitext(file.filename)[1]
    input_path = os.path.join(UPLOAD_DIR, unique_id + ext)
    output_path = get_unique_filename(OUTPUT_DIR, f"{unique_id}.pdf")

    contents = await file.read()
    with open(input_path, "wb") as f:
        f.write(contents)

    try:
        excel_to_pdf(input_path, output_path)
    except Exception as e:
        return {"error": str(e)}

    threading.Thread(target=delete_later, args=(input_path,)).start()
    threading.Thread(target=delete_later, args=(output_path,)).start()

    return FileResponse(output_path, media_type="application/pdf", filename=f"{os.path.splitext(file.filename)[0]}.pdf")


# ---------------- PPT to PDF ----------------
@app.post("/convert/ppt-to-pdf")
async def convert_ppt(file: UploadFile = File(...)):
    if not file.filename.lower().endswith((".pptx", ".ppt")):
        return {"error": "Only PowerPoint files allowed"}

    unique_id = str(uuid.uuid4())
    ext = os.path.splitext(file.filename)[1]
    input_path = os.path.join(UPLOAD_DIR, unique_id + ext)
    contents = await file.read()
    with open(input_path, "wb") as f:
        f.write(contents)

    try:
        subprocess.run([
            "soffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", OUTPUT_DIR,
            input_path
        ], check=True)
    except Exception as e:
        return {"error": str(e)}

    output_file = os.path.join(OUTPUT_DIR, os.path.splitext(os.path.basename(input_path))[0] + ".pdf")
    threading.Thread(target=delete_later, args=(input_path,)).start()
    threading.Thread(target=delete_later, args=(output_file,)).start()

    return FileResponse(output_file, media_type="application/pdf", filename=f"{os.path.splitext(file.filename)[0]}.pdf")


# ---------------- Run Server ----------------
if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)